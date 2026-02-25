"""
Export Service – Multi-Format Document Generation
──────────────────────────────────────────────────
Converts verified Markdown text into PDF, DOCX, or raw Markdown files.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Optional

logger = logging.getLogger(__name__)


def _md_to_plain(md_text: str) -> str:
    """Strip basic Markdown formatting for plain-text export."""
    text = re.sub(r"#{1,6}\s*", "", md_text)  # headings
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # bold
    text = re.sub(r"\*(.+?)\*", r"\1", text)  # italic
    text = re.sub(r"_(.+?)_", r"\1", text)  # italic
    text = re.sub(r"`(.+?)`", r"\1", text)  # inline code
    text = re.sub(r"^\s*[-*+]\s", "- ", text, flags=re.MULTILINE)  # bullets
    return text.strip()


def export_markdown(text: str, title: str = "OmniDraft Note") -> bytes:
    """Return raw Markdown bytes."""
    header = f"# {title}\n\n"
    return (header + text).encode("utf-8")


def export_txt(text: str, title: str = "OmniDraft Note") -> bytes:
    """Return plain text bytes."""
    plain = _md_to_plain(text)
    header = f"{title}\n{'=' * len(title)}\n\n"
    return (header + plain).encode("utf-8")


def export_pdf(text: str, title: str = "OmniDraft Note") -> bytes:
    """Generate a styled PDF from the transcription text."""
    from fpdf import FPDF
    import os
    
    # Resolve the path to the fonts directory relative to this service file
    current_dir = os.path.dirname(os.path.abspath(__file__))
    fonts_dir = os.path.join(current_dir, "..", "fonts")
    
    font_regular = os.path.join(fonts_dir, "NotoSans-Regular.ttf")
    font_bold = os.path.join(fonts_dir, "NotoSans-Bold.ttf")
    font_telugu = os.path.join(fonts_dir, "NotoSansTelugu-Regular.ttf")

    class ScribePDF(FPDF):
        def header(self):
            # Using NotoSans for header
            self.set_font("NotoSans", "B", 10)
            self.set_text_color(100, 100, 100)
            self.cell(0, 8, "OmniDraft", align="R")
            self.ln(12)

        def footer(self):
            self.set_y(-15)
            self.set_font("NotoSans", "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    pdf = ScribePDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Add Unicode fonts
    pdf.add_font("NotoSans", "", font_regular)
    pdf.add_font("NotoSans", "B", font_bold)
    pdf.add_font("NotoSans", "I", font_regular) # fallback for italic
    pdf.add_font("NotoSansTelugu", "", font_telugu)
    
    # Set fallback fonts for Unicode support (e.g. Telugu)
    pdf.set_fallback_fonts(["NotoSansTelugu"])
    
    # Enable text shaping for correct rendering of complex scripts (requires uharfbuzz)
    pdf.set_text_shaping(True)
    
    pdf.add_page()

    # Title
    pdf.set_font("NotoSans", "B", 20)
    pdf.set_text_color(30, 30, 30)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Divider line
    pdf.set_draw_color(70, 130, 180)
    pdf.set_line_width(0.8)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(8)

    # Body text – construct formatted markdown elements
    # Using `text` directly (skip _md_to_plain) so we can style headings and bold text
    pdf.set_font("NotoSans", "", 11)
    pdf.set_text_color(40, 40, 40)

    # Body text – construct formatted markdown elements logically by paragraph/block
    blocks = re.split(r'\n\s*\n', text.strip())
    
    for block in blocks:
        block = block.strip()
        if not block:
            continue
            
        # 1. Is this a Heading?
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", block, flags=re.MULTILINE|re.DOTALL)
        if heading_match and "\n" not in heading_match.group(2):
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).replace("\n", " ").strip()
            pdf.ln(4)
            size = max(11, 20 - (level * 2))
            pdf.set_font("NotoSans", "B", size)
            pdf.set_text_color(20, 20, 20)
            pdf.multi_cell(0, size * 0.4, heading_text, new_x="LMARGIN", new_y="NEXT", markdown=True)
            pdf.ln(2)
            continue
            
        # 2. Is this a Table block?
        if "\n" in block and "|" in block.split("\n")[0] and "|" in block.split("\n")[1]:
            lines = block.split("\n")
            pdf.ln(2)
            pdf.set_font("NotoSans", "", 10)
            
            # Use fpdf2 built-in table rendering
            with pdf.table(text_align="LEFT") as table:
                for i, line in enumerate(lines):
                    # Skip the markdown separator row (e.g. |---|---|)
                    if set(line.strip().replace(" ", "").replace("|", "").replace("-", "")) == set():
                        continue
                    
                    # Parse row cells
                    cells = [c.strip() for c in line.strip().strip("|").split("|")]
                    row_ctx = table.row()
                    for cell in cells:
                        if i == 0:
                            pdf.set_font("NotoSans", "B", 10)
                        else:
                            pdf.set_font("NotoSans", "", 10)
                        pdf.set_text_color(40, 40, 40)
                        row_ctx.cell(cell)
            pdf.ln(4)
            continue
            
        # 3. Is this a Code block?
        if block.startswith("```") and block.endswith("```"):
            code_content = block[3:-3].strip()
            if "\n" in code_content:
                first_line, rest = code_content.split("\n", 1)
                if not " " in first_line:
                    code_content = rest
            
            pdf.ln(2)
            pdf.set_font("Courier", "", 10)
            pdf.set_text_color(60, 60, 60)
            pdf.set_fill_color(245, 245, 245)
            for line in code_content.split("\n"):
                pdf.multi_cell(0, 5, line, new_x="LMARGIN", new_y="NEXT", markdown=False, fill=True)
            pdf.ln(2)
            continue
            
        # 3. Is this a List block? (Bullet / Numbered)
        if block.startswith("- ") or block.startswith("* ") or re.match(r"^\d+\.\s+", block):
            pdf.ln(2)
            lines = block.split("\n")
            for line in lines:
                line = line.strip()
                if not line: continue
                # Bullet
                if line.startswith("- ") or line.startswith("* "):
                    content = line[2:].strip()
                    pdf.set_font("NotoSans", "", 11)
                    pdf.set_text_color(40, 40, 40)
                    pdf.set_x(15)
                    pdf.multi_cell(0, 6, f"\u2022 {content}", new_x="LMARGIN", new_y="NEXT", markdown=True)
                # Numbered
                elif re.match(r"^(\d+\.)\s+(.*)", line):
                    match = re.match(r"^(\d+\.)\s+(.*)", line)
                    pdf.set_font("NotoSans", "", 11)
                    pdf.set_text_color(40, 40, 40)
                    pdf.set_x(15)
                    pdf.multi_cell(0, 6, f"{match.group(1)} {match.group(2).strip()}", new_x="LMARGIN", new_y="NEXT", markdown=True)
                else:
                    # indented continuation line
                    pdf.set_font("NotoSans", "", 11)
                    pdf.set_text_color(40, 40, 40)
                    pdf.set_x(20)
                    pdf.multi_cell(0, 6, line, new_x="LMARGIN", new_y="NEXT", markdown=True)
            pdf.ln(2)
            continue
            
        # 4. Regular paragraph - replace internal newlines with spaces so fpdf2 wraps beautifully and naturally
        paragraph_text = block.replace("\n", " ").strip()
        pdf.set_font("NotoSans", "", 11)
        pdf.set_text_color(40, 40, 40)
        pdf.set_x(10)
        pdf.multi_cell(0, 6, paragraph_text, new_x="LMARGIN", new_y="NEXT", markdown=True)
        pdf.ln(3)

    return bytes(pdf.output())


def export_docx(text: str, title: str = "OmniDraft Note") -> bytes:
    """Generate a styled DOCX from the transcription text."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Style the title
    title_para = doc.add_heading(title, level=0)
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(30, 30, 30)

    # Add a subtle horizontal rule
    rule_para = doc.add_paragraph()
    rule_para.paragraph_format.space_after = Pt(12)
    rule_run = rule_para.add_run("─" * 60)
    rule_run.font.color.rgb = RGBColor(70, 130, 180)
    rule_run.font.size = Pt(8)

    # Body text – construct formatted markdown elements logically by paragraph/block
    blocks = re.split(r'\n\s*\n', text.strip())
    
    for block in blocks:
        block = block.strip()
        if not block:
            continue

        # 1. Headings
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", block, flags=re.MULTILINE|re.DOTALL)
        if heading_match and "\n" not in heading_match.group(2):
            level = min(len(heading_match.group(1)), 4)
            heading_text = heading_match.group(2).replace("\n", " ").strip()
            doc.add_heading(heading_text, level=level)
            continue

        # 2. Tables
        if "\n" in block and "|" in block.split("\n")[0] and "|" in block.split("\n")[1]:
            lines = block.split("\n")
            # find number of columns from first row
            header_cells = [c.strip() for c in lines[0].strip().strip("|").split("|")]
            table = doc.add_table(rows=0, cols=len(header_cells))
            table.style = 'Table Grid'
            
            for i, line in enumerate(lines):
                # Skip the markdown separator row (e.g. |---|---|)
                if set(line.strip().replace(" ", "").replace("|", "").replace("-", "")) == set():
                    continue
                
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                row_cells = table.add_row().cells
                for j, cell in enumerate(cells):
                    if j < len(row_cells):
                        row_cells[j].text = cell
                        if i == 0:  # Bold headers
                            for paragraph in row_cells[j].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
            doc.add_paragraph("") # Space after table
            continue

        # 3. Code blocks
        if block.startswith("```") and block.endswith("```"):
            code_content = block[3:-3].strip()
            if "\n" in code_content:
                first_line, rest = code_content.split("\n", 1)
                if not " " in first_line:
                    code_content = rest
            
            para = doc.add_paragraph(code_content)
            for run in para.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(80, 80, 80)
            continue

        # 4. List block (Bullet / Numbered)
        if block.startswith("- ") or block.startswith("* ") or re.match(r"^\d+\.\s+", block):
            lines = block.split("\n")
            for line in lines:
                line = line.strip()
                if not line: continue
                # Bullet
                if line.startswith("- ") or line.startswith("* "):
                    content = line[2:].strip()
                    doc.add_paragraph(content, style="List Bullet")
                # Numbered
                elif re.match(r"^(\d+\.)\s+(.*)", line):
                    match = re.match(r"^(\d+\.)\s+(.*)", line)
                    doc.add_paragraph(match.group(2).strip(), style="List Number")
                else:
                    doc.add_paragraph(line)
            continue

        # 5. Regular paragraph
        paragraph_text = block.replace("\n", " ").strip()
        para = doc.add_paragraph()
        
        # Handle bold and italic in-line
        segments = re.split(r"(\*\*.*?\*\*|\*.*?\*|_.*?_|`.*?`)", paragraph_text)
        for seg in segments:
            if seg.startswith("**") and seg.endswith("**"):
                run = para.add_run(seg[2:-2])
                run.bold = True
            elif (seg.startswith("*") and seg.endswith("*")) or \
                 (seg.startswith("_") and seg.endswith("_")):
                run = para.add_run(seg[1:-1])
                run.italic = True
            elif seg.startswith("`") and seg.endswith("`"):
                run = para.add_run(seg[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(180, 60, 60)
            else:
                para.add_run(seg)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_note(
    text: str,
    format: str,
    title: str = "OmniDraft Note",
) -> tuple[bytes, str, str]:
    """
    Export a note to the requested format.

    Returns (file_bytes, content_type, file_extension).
    """
    format_lower = format.lower().strip(".")

    exporters = {
        "pdf": (export_pdf, "application/pdf", ".pdf"),
        "docx": (export_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx"),
        "md": (export_markdown, "text/markdown", ".md"),
        "markdown": (export_markdown, "text/markdown", ".md"),
        "txt": (export_txt, "text/plain", ".txt"),
    }

    if format_lower not in exporters:
        raise ValueError(f"Unsupported format: {format}. Supported: {list(exporters.keys())}")

    func, content_type, ext = exporters[format_lower]
    file_bytes = func(text, title)
    return file_bytes, content_type, ext
